#!/usr/bin/env python3
"""
Markdownレポートを .docx に変換するスクリプト
research_team.py と連携して使用
"""

from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def md_to_docx(md_path: str, docx_path: Optional[str] = None):
    """Markdown ファイルを Word 文書に変換"""
    md_path = Path(md_path)
    if not md_path.exists():
        print(f"Error: {md_path} not found")
        sys.exit(1)

    if docx_path is None:
        docx_path = md_path.with_suffix(".docx")
    else:
        docx_path = Path(docx_path)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # -- Style setup --
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # -- Parse markdown line by line --
    i = 0
    table_rows = []  # Buffer for table parsing

    while i < len(lines):
        line = lines[i]

        # --- Table detection ---
        if "|" in line and line.strip().startswith("|"):
            # Collect all table lines
            table_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                # Skip separator rows (|---|---|)
                if not re.match(r"^\|[\s\-:|]+\|$", row):
                    cells = [c.strip() for c in row.split("|")[1:-1]]
                    table_rows.append(cells)
                i += 1
            _add_table(doc, table_rows)
            doc.add_paragraph()  # Space after table
            continue

        # --- Headings ---
        if line.startswith("# "):
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)

        # --- Horizontal rule ---
        elif line.strip() == "---":
            pass  # Skip, headings provide structure

        # --- Blockquote ---
        elif line.strip().startswith("> "):
            quote_text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # --- Bold metadata lines (e.g. **作成日：**...) ---
        elif line.startswith("**") and "**" in line[2:]:
            p = doc.add_paragraph()
            _add_rich_text(p, line)

        # --- Bullet list ---
        elif line.strip().startswith("- "):
            content = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, content)

        # --- Numbered list ---
        elif re.match(r"^\d+\.\s", line.strip()):
            content = re.sub(r"^\d+\.\s", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, content)

        # --- Empty line ---
        elif line.strip() == "":
            pass  # Skip blank lines (spacing handled by paragraph format)

        # --- Regular paragraph ---
        else:
            # Collect consecutive non-empty, non-special lines as one paragraph
            para_lines = [line]
            while (i + 1 < len(lines)
                   and lines[i + 1].strip() != ""
                   and not lines[i + 1].startswith("#")
                   and not lines[i + 1].startswith("|")
                   and not lines[i + 1].startswith(">")
                   and not lines[i + 1].startswith("- ")
                   and not re.match(r"^\d+\.\s", lines[i + 1].strip())
                   and not lines[i + 1].strip().startswith("**")
                   and lines[i + 1].strip() != "---"):
                i += 1
                para_lines.append(lines[i])

            full_text = " ".join(l.strip() for l in para_lines)
            if full_text.strip():
                p = doc.add_paragraph()
                _add_rich_text(p, full_text)

        i += 1

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")
    return docx_path


def _add_rich_text(paragraph, text: str):
    """Markdown の太字 (**text**) と引用 ([n]) をリッチテキストに変換"""
    # Split by **bold** markers
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle citation markers [1], [2] etc
            sub_parts = re.split(r"(\[\d+\])", part)
            for sp in sub_parts:
                if re.match(r"\[\d+\]", sp):
                    run = paragraph.add_run(sp)
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x2D, 0x5A, 0x27)
                else:
                    paragraph.add_run(sp)


def _add_table(doc, rows: List[List[str]]):
    """表を Word 文書に追加"""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = cell_text
                # Bold the header row
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 convert_report.py <markdown_file> [output.docx]")
        sys.exit(1)

    md_file = sys.argv[1]
    out_file = sys.argv[2] if len(sys.argv) > 2 else None
    md_to_docx(md_file, out_file)
